import os.path
import numpy as np
import pymupdf
import re
import pickle
import docx
import Report
import docx2pdf

class ClientProfile:
    def __init__(self, clientName):
        self.clientName = clientName
        self.reports = []
        self.savePath = None
        self.clientFolders = []
        self.supportedFileFormats = [".pdf",".doc",".docx"]

    def AddReport(self, report):
        self.reports.append(report)
        report.ExtractText()

    def SearchReports(self, folder):
        if folder not in self.clientFolders:
            self.clientFolders.append(folder)
        reports = []
        contents = os.listdir(folder)
        dirs = [_dir for _dir in contents if os.path.isdir(folder+"\\"+_dir)]
        reportWords = ["report", "final","rep", "hlc"]

        for _object in contents:
            if"." + _object.split(".")[-1].lower() not in self.supportedFileFormats:
                continue
            for reportWord in reportWords:
                if reportWord in _object.lower():
                    path = folder+"\\"+_object
                    if path not in reports and path not in dirs:
                        print(path)
                        self.AddReport(Report(path))
                        break
        for _dir in dirs:
            self.SearchReports(folder+"\\"+_dir)
        print(self.reports)

    def GenerateExecSumDoc(self):
        summaryDoc = docx.Document()
        savePath = f"{self.clientFolders[0]}\\ExecutiveSummaryReport.docx"
        reports = self.reports
        reports = sorted(reports, key = lambda x: x.date[0] if len(x.date)>0 else "")
        print(reports)
        for report in reports:

            if len(report.execSum)==0:
                continue

            print(title + "\n\n" for title in report.title)
            try:
                summaryDoc.add_heading(report.title[0],2)
            except Exception as e:
                print(f"Failed to add title to document, following error occured:\n{e}")

            try:
                summaryDoc.add_paragraph(text = f"{report.date[0]} \n" )
            except Exception as e:
                print(f"Failed to add Date/Authours to document, following error occured:\n{e}")

            try:
                summaryDoc.add_paragraph(text = f"By {', '.join([author for author in report.authours])}\n")
            except Exception as e:
                print(f"Failed to add Date/Authours to document, following error occured:\n{e}")

            try:
                summaryDoc.add_paragraph(text = report.execSum[0] + "\n")
            except Exception as e:
                print(f"Failed to add summary to document, following error occured:\n{e}")

            summaryDoc.add_page_break()
        summaryDoc.save(savePath)
        print(f"Saved to {savePath}")


def Main():
    myClient = ClientProfile("MillMerran")

    myClient.clientFolders.append("L:\hrltec\OPS\Millmerran")
    myClient.SearchReports("L:\hrltec\OPS\Millmerran")

    # myReports = [
    #     Report(r"F:\Python\Scripts\Searches\Rep_HLC2017149-M7.pdf"),
    #     Report(r"L:\hrltec\OPS\Millmerran\50115100 - Final Reheater Tubing Creep-Rupture Testing\Reporting\Rep_HLC2013007.pdf"),
    #     Report(r"L:\hrltec\OPS\Millmerran\48231911 - Rotary Air Heater Engineering Assessment\Reporting\HRL HLC2023127 MOC RAH Engineering Assessment R1.pdf"),
    #     Report(r"L:\hrltec\OPS\Millmerran\48221881 - MOC N1 and N2 linked piping to HP8 heater Flexibility Analysis\HRL Report HLC_2022_135 MOC_Unit 1 HP8 Heater FEA and DTA.pdf"),
    #     Report(r"L:\hrltec\OPS\Melbourne Water\50215621 - Material verification for valve replacement project\Reporting\50215621 Final Report.pdf"),
    #     Report(r"L:\hrltec\OPS\Millmerran\48221869 - Millmerran Defect Tolerance Assessment\hrl hlcreport QA720I_4822135_MOC_HP8 Heater_DTA_D3.docm")]

    # for report in myReports:
    #     myClient.AddReport(report)
    # for rep in myClient.reports:
    #     for k, v in vars(rep).items():
    #         if not (k == "body"):
    #             print(k, v)

    myClient.GenerateExecSumDoc()




if __name__ == '__main__':
    Main()





