from __future__ import annotations

import argparse
import csv
import json
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from openpyxl import Workbook

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from client_profiler import ClientProfiler, ProfilerConfig


DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "sample_docs" / "financial_variants"
DEFAULT_EXPECTED_PATH = PROJECT_ROOT / "data" / "financial_variants_expected.json"


@dataclass
class FinancialDoc:
    client: str
    project: str
    project_code: str
    filename: str
    kind_hint: str
    body: dict[str, Any]


def _documents() -> list[FinancialDoc]:
    return [
        FinancialDoc(
            client="NorthRiver Energy",
            project="Turbine Outage Stabilization",
            project_code="NRE-TUR-OUT-2026",
            filename="invoice_nre_turbine_outage_2026.txt",
            kind_hint="invoice",
            body={
                "Client Name": "NorthRiver Energy",
                "Project": "Turbine Outage Stabilization",
                "Project Code": "NRE-TUR-OUT-2026",
                "Date": "2026-04-12",
                "Invoice Number": "INV-NRE-260412",
                "Invoice Total": "USD 182000.00",
                "Cost": "USD 117500.00",
                "Notes": "Includes emergency callout standby and overtime labor.",
            },
        ),
        FinancialDoc(
            client="NorthRiver Energy",
            project="Turbine Outage Stabilization",
            project_code="NRE-TUR-OUT-2026",
            filename="expense_report_nre_travel_and_lift.md",
            kind_hint="expense_report",
            body={
                "Client": "NorthRiver Energy",
                "Project": "Turbine Outage Stabilization",
                "Project Code": "NRE-TUR-OUT-2026",
                "Date": "2026-04-10",
                "Expense Report": "Travel and lifting equipment",
                "Travel Cost": "USD 13600.00",
                "Material Cost": "USD 4300.00",
                "Total Cost": "USD 17900.00",
            },
        ),
        FinancialDoc(
            client="BlueMesa Utilities",
            project="Switchyard Reliability Program",
            project_code="BMU-SWY-REL-2026",
            filename="quote_bmu_switchyard_program.html",
            kind_hint="quote",
            body={
                "Client Name": "BlueMesa Utilities",
                "Project": "Switchyard Reliability Program",
                "Project Code": "BMU-SWY-REL-2026",
                "Date": "2026-03-22",
                "Quote Number": "Q-BMU-0322",
                "Contract Value": "USD 96500.00",
                "Estimated Cost": "USD 60200.00",
                "Scope": "Relay verification, outage sequencing, close-out reporting",
            },
        ),
        FinancialDoc(
            client="BlueMesa Utilities",
            project="Switchyard Reliability Program",
            project_code="BMU-SWY-REL-2026",
            filename="timesheet_bmu_night_shift_support.csv",
            kind_hint="timesheet",
            body={
                "Client": "BlueMesa Utilities",
                "Project": "Switchyard Reliability Program",
                "Project Code": "BMU-SWY-REL-2026",
                "Date": "2026-03-18",
                "Timesheet": "Night shift support",
                "Hours Worked": "168",
                "Labour Cost": "USD 15120.00",
                "Expense": "USD 980.00",
                "Total Cost": "USD 16100.00",
            },
        ),
        FinancialDoc(
            client="HarborStone Manufacturing",
            project="Compressor Reliability Reset",
            project_code="HSM-COM-REL-2026",
            filename="cost_breakdown_hsm_compressor_reset.xlsx",
            kind_hint="expense_report",
            body={
                "Client": "HarborStone Manufacturing",
                "Project": "Compressor Reliability Reset",
                "Project Code": "HSM-COM-REL-2026",
                "Date": "2026-02-09",
                "Cost Breakdown": "Parts, labor, and access",
                "Material Cost": "USD 28750.00",
                "Labour Cost": "USD 33400.00",
                "Travel Cost": "USD 5100.00",
                "Total Cost": "USD 67250.00",
            },
        ),
        FinancialDoc(
            client="HarborStone Manufacturing",
            project="Compressor Reliability Reset",
            project_code="HSM-COM-REL-2026",
            filename="invoice_hsm_compressor_closeout.docx",
            kind_hint="invoice",
            body={
                "Client Name": "HarborStone Manufacturing",
                "Project": "Compressor Reliability Reset",
                "Project Code": "HSM-COM-REL-2026",
                "Date": "2026-02-28",
                "Invoice Number": "INV-HSM-0228",
                "Invoice Total": "USD 121300.00",
                "Cost": "USD 74420.00",
                "Gross Profit": "USD 46880.00",
                "Approval": "Approved by Plant Reliability Manager",
            },
        ),
        FinancialDoc(
            client="Delta Utilities",
            project="Waterline Integrity Sprint",
            project_code="DEL-WAT-INT-2026",
            filename="procurement_summary_delta_waterline.txt",
            kind_hint="purchase_order",
            body={
                "Client": "Delta Utilities",
                "Project": "Waterline Integrity Sprint",
                "Project Code": "DEL-WAT-INT-2026",
                "Date": "2026-05-04",
                "Purchase Order Number": "PO-DEL-0504",
                "Supplier Cost": "USD 33980.00",
                "Expense": "USD 4200.00",
                "Total Cost": "USD 38180.00",
                "Comment": "Direct reimbursable procurement package",
            },
        ),
        FinancialDoc(
            client="Delta Utilities",
            project="Waterline Integrity Sprint",
            project_code="DEL-WAT-INT-2026",
            filename="change_order_delta_waterline.md",
            kind_hint="quote",
            body={
                "Client Name": "Delta Utilities",
                "Project": "Waterline Integrity Sprint",
                "Project Code": "DEL-WAT-INT-2026",
                "Date": "2026-05-12",
                "Quote Number": "Q-DEL-0512",
                "Contract Value": "USD 58900.00",
                "Estimated Cost": "USD 34870.00",
                "Reason": "Additional hydro test and trench restoration",
            },
        ),
        FinancialDoc(
            client="BlueMesa Utilities",
            project="Switchyard Reliability Program",
            project_code="BMU-SWY-REL-2026",
            filename="narrative_profit_note_bmu.html",
            kind_hint="report",
            body={
                "Client": "BlueMesa Utilities",
                "Project": "Switchyard Reliability Program",
                "Project Code": "BMU-SWY-REL-2026",
                "Date": "2026-03-31",
                "Revenue": "USD 12500.00",
                "Cost": "USD 6400.00",
                "Summary": "Interim progress claim for completed relay acceptance tests",
            },
        ),
    ]


def _expected_totals(docs: list[FinancialDoc]) -> dict[str, dict[str, float]]:
    totals: dict[str, dict[str, float]] = {}
    for doc in docs:
        entry = totals.setdefault(doc.client, {"revenue": 0.0, "cost": 0.0})
        revenue = _parse_money(doc.body.get("Invoice Total") or doc.body.get("Contract Value") or doc.body.get("Revenue"))
        cost = _parse_money(
            doc.body.get("Total Cost")
            or doc.body.get("Estimated Cost")
            or doc.body.get("Cost")
            or doc.body.get("Labour Cost")
        )
        entry["revenue"] += revenue
        entry["cost"] += cost

    for client, row in totals.items():
        row["revenue"] = round(row["revenue"], 2)
        row["cost"] = round(row["cost"], 2)
        row["profit"] = round(row["revenue"] - row["cost"], 2)
    return totals


def _parse_money(value: Any) -> float:
    if value is None:
        return 0.0
    text = str(value).strip().replace(",", "")
    if not text:
        return 0.0
    cleaned = "".join(ch for ch in text if ch.isdigit() or ch in {".", "-"})
    try:
        return float(cleaned) if cleaned else 0.0
    except ValueError:
        return 0.0


def _write_doc(path: Path, payload: FinancialDoc) -> None:
    suffix = path.suffix.lower()
    data = payload.body

    if suffix == ".txt":
        lines = [f"{payload.project} Financial Document", ""]
        for k, v in data.items():
            lines.append(f"{k}: {v}")
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    if suffix in {".md", ".markdown"}:
        lines = [f"# {payload.project} Financial Update", ""]
        for k, v in data.items():
            lines.append(f"- **{k}:** {v}")
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    if suffix in {".html", ".htm"}:
        rows = "".join(f"<tr><th>{k}</th><td>{v}</td></tr>" for k, v in data.items())
        html = (
            "<html><body>"
            f"<h1>{payload.project} Commercial Snapshot</h1>"
            f"<table border='1'>{rows}</table>"
            "</body></html>"
        )
        path.write_text(html, encoding="utf-8")
        return

    if suffix == ".csv":
        with path.open("w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["field", "value"])
            for k, v in data.items():
                writer.writerow([k, v])
        return

    if suffix == ".xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = "financial"
        ws.append(["field", "value"])
        for k, v in data.items():
            ws.append([k, v])
        wb.save(path)
        return

    if suffix == ".docx":
        doc = DocxDocument()
        doc.add_heading(f"{payload.project} Commercial Document", level=1)
        for k, v in data.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.save(path)
        return

    raise ValueError(f"Unsupported suffix for generation: {suffix}")


def _write_documents(output_dir: Path, docs: list[FinancialDoc]) -> list[dict[str, Any]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: list[dict[str, Any]] = []
    for item in docs:
        path = output_dir / item.filename
        _write_doc(path, item)
        manifest.append(
            {
                "path": str(path),
                "client": item.client,
                "project": item.project,
                "project_code": item.project_code,
                "kind_hint": item.kind_hint,
            }
        )
    return manifest


def _ingest(output_dir: Path, db_path: Path, force: bool) -> list[dict[str, Any]]:
    config = ProfilerConfig(db_path=db_path, data_dir=db_path.parent, llm_provider="none")
    profiler = ClientProfiler(config)
    return profiler.ingest_directory(output_dir, force_reingest=force)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate varied fictional financial documents for client profiling")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--expected-out", type=Path, default=DEFAULT_EXPECTED_PATH)
    parser.add_argument("--ingest", action="store_true")
    parser.add_argument("--db", type=Path, default=Path("./data/profiler_financial_variants.db"))
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()

    docs = _documents()
    manifest = _write_documents(args.output_dir, docs)
    expected = _expected_totals(docs)

    args.expected_out.parent.mkdir(parents=True, exist_ok=True)
    args.expected_out.write_text(json.dumps(expected, indent=2), encoding="utf-8")

    result: dict[str, Any] = {
        "generated_documents": len(manifest),
        "output_dir": str(args.output_dir),
        "expected_out": str(args.expected_out),
        "expected_totals": expected,
    }

    if args.ingest:
        ingest_rows = _ingest(args.output_dir, args.db, args.force)
        result["ingested"] = len(ingest_rows)
        result["ingest_status_counts"] = {
            key: sum(1 for row in ingest_rows if row.get("status") == key)
            for key in sorted({str(row.get("status") or "") for row in ingest_rows})
        }

    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
