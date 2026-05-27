from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


@dataclass
class ReportSpec:
    year: int
    client: str
    job_type: str
    fmt: str
    site: str
    asset: str
    report_id: str
    author: str
    contact: str


def section_paragraphs(spec: ReportSpec) -> dict[str, list[str]]:
    return {
        "Introduction": [
            (
                f"This report documents a {spec.job_type.lower()} conducted for {spec.client} at {spec.site} during {spec.year}. "
                "The engagement was initiated after a sequence of repeat defects and growing maintenance backlog signaled increased reliability risk. "
                "The client requested an evidence-led assessment with practical recommendations that could be executed within routine shutdown constraints."
            ),
            (
                "The study is written for operations management, maintenance planners, and integrity engineers. "
                "It consolidates field observations, historical maintenance trends, and fit-for-purpose risk ranking so that decisions on intervention timing, "
                "resource allocation, and assurance controls can be made with clear traceability."
            ),
        ],
        "Scope": [
            (
                f"Scope covered the nominated asset boundary around {spec.asset}, including adjacent interfaces that materially influence reliability performance. "
                "Activities included document review, condition walkdowns, selected checks against design intent, and stakeholder interviews across operations and maintenance shifts."
            ),
            (
                "The assessment deliberately excluded full detail design and procurement packaging. "
                "However, it included enough technical depth to define immediate corrective actions, medium-term mitigation options, and an implementation sequence "
                "that can be integrated into existing planning cycles."
            ),
        ],
        "Background": [
            (
                f"Between {spec.year - 3} and {spec.year - 1}, work order data indicated increasing corrective intervention frequency against the assessed equipment class. "
                "Temporary repairs and repeat call-outs were more common where close-out evidence was incomplete or acceptance criteria were interpreted inconsistently."
            ),
            (
                "Asset criticality workshops previously identified this area as contributing disproportionately to production interruptions. "
                "Given the current cost pressure and constrained outage windows, the client sought a targeted approach that could stabilize performance "
                "without requiring immediate major capital replacement."
            ),
        ],
        "Method/Work Done": [
            (
                "The team applied a four-stage method: evidence collection, condition grading, risk ranking, and stakeholder validation. "
                "Evidence collection combined maintenance history, calibration/test records, and field observations captured against a standardized checklist."
            ),
            (
                "Condition grading used a weighted framework incorporating severity, recurrence, and detectability. "
                "Findings were reviewed in a multidisciplinary workshop to test assumptions, validate practicality, and assign accountable owners for each action item."
            ),
        ],
        "Results": [
            (
                "The investigation identified a pattern of moderate-to-high risk issues concentrated in interfaces between routine maintenance and verification workflows. "
                "While several defects were technical in nature, the stronger predictor of recurrence was inconsistency in post-work validation evidence."
            ),
            (
                "Performance trend analysis over twelve months indicated that assets with explicit acceptance criteria and documented verification closed with lower repeat failure rates. "
                "The included chart and supporting table summarise monthly defect counts and mean time between failures for the assessed period."
            ),
        ],
        "Discussion": [
            (
                "The results suggest that standalone repairs are unlikely to deliver durable improvement unless paired with process controls. "
                "In particular, inadequate handover detail between field execution and reliability assurance allows latent defects to persist and re-emerge under load."
            ),
            (
                "A balanced strategy is therefore recommended: immediate defect elimination for highest-risk items, combined with stronger verification standards, "
                "clear acceptance criteria, and periodic governance reviews to ensure corrective intent is sustained in practice."
            ),
        ],
        "Conclusion": [
            (
                f"Overall condition for {spec.asset} is serviceable but vulnerable to escalating reliability risk if current execution variability continues. "
                "The asset can remain in operation provided prioritized interventions are completed and assurance controls are strengthened."
            ),
            (
                "The assessment confirms there is no single root cause; rather, risk is generated by the interaction of asset degradation, planning constraints, "
                "and inconsistent verification discipline. A staged improvement program over two quarters is expected to materially reduce reactive demand."
            ),
        ],
        "Recommendations": [
            (
                "1) Complete high-priority corrective actions within 30 days and capture objective close-out evidence. "
                "2) Standardize maintenance work packs with explicit acceptance criteria and mandatory verification checkpoints."
            ),
            (
                "3) Establish monthly reliability review using defect aging, recurrence, and verification quality KPIs. "
                "4) Perform a six-month reassessment to confirm risk reduction and identify residual gaps requiring further intervention."
            ),
        ],
    }


def build_series(seed: int) -> tuple[list[str], list[int], list[float]]:
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    defects = [max(2, (seed * 3 + i * 2) % 11 + 2) for i in range(12)]
    mtbf = [round(120 - d * 4.3 + (i % 3) * 2.1, 1) for i, d in enumerate(defects)]
    return months, defects, mtbf


def create_chart(chart_path: Path, title: str, seed: int) -> tuple[list[str], list[int], list[float]]:
    months, defects, mtbf = build_series(seed)
    fig, ax1 = plt.subplots(figsize=(8.0, 3.8))
    ax1.plot(months, defects, marker="o", linewidth=2.0, color="#c0392b", label="Defect Count")
    ax1.set_ylabel("Defects / month", color="#c0392b")
    ax1.tick_params(axis="y", labelcolor="#c0392b")
    ax1.grid(True, axis="y", linestyle="--", alpha=0.4)

    ax2 = ax1.twinx()
    ax2.plot(months, mtbf, marker="s", linewidth=2.0, color="#1f618d", label="MTBF (hours)")
    ax2.set_ylabel("MTBF (h)", color="#1f618d")
    ax2.tick_params(axis="y", labelcolor="#1f618d")

    fig.suptitle(title, fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(chart_path, dpi=180)
    plt.close(fig)
    return months, defects, mtbf


def add_docx_header_footer(doc: Document, letterhead: str) -> None:
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = letterhead
    sec.header.paragraphs[0].runs[0].bold = True
    sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    sec.footer.paragraphs[0].text = "Confidential - Consultancy Technical Report"
    sec.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_docx(out: Path, spec: ReportSpec, sections: dict[str, list[str]], chart_path: Path, months: list[str], defects: list[int], mtbf: list[float]) -> None:
    doc = Document()
    add_docx_header_footer(doc, "ALINTA CONSULTING SERVICES")

    title = doc.add_paragraph()
    run = title.add_run("TECHNICAL CONSULTING REPORT")
    run.bold = True
    run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Client Name: {spec.client}\n").bold = True
    meta.add_run(f"Report ID: {spec.report_id}\n")
    meta.add_run(f"Report Date: {spec.year}-10-15\n")
    meta.add_run(f"Client Contact: {spec.contact}\n")
    meta.add_run(f"Author: {spec.author}\n")
    meta.add_run(f"Site: {spec.site}\nAsset: {spec.asset}")

    doc.add_page_break()

    for name, paras in sections.items():
        h = doc.add_paragraph()
        r = h.add_run(name)
        r.bold = True
        r.font.size = Pt(16)
        for p in paras:
            para = doc.add_paragraph(p)
            para.paragraph_format.space_after = Pt(8)

        if name == "Results":
            doc.add_paragraph("Results Data Trend Chart").runs[0].bold = True
            doc.add_picture(str(chart_path), width=Inches(6.4))
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Month"
            hdr[1].text = "Defects"
            hdr[2].text = "MTBF (h)"
            for m, d, t in zip(months, defects, mtbf):
                row = table.add_row().cells
                row[0].text = m
                row[1].text = str(d)
                row[2].text = str(t)

    doc.save(out)


def write_html(out: Path, spec: ReportSpec, sections: dict[str, list[str]], chart_path: Path, months: list[str], defects: list[int], mtbf: list[float]) -> None:
    chart_name = chart_path.name
    section_html = []
    for name, paras in sections.items():
        section_html.append(f"<h2>{name}</h2>")
        section_html.extend([f"<p>{p}</p>" for p in paras])
        if name == "Results":
            rows = "".join(
                f"<tr><td>{m}</td><td>{d}</td><td>{t}</td></tr>" for m, d, t in zip(months, defects, mtbf)
            )
            section_html.append(f"<img src=\"{chart_name}\" alt=\"Trend chart\" class=\"chart\" />")
            section_html.append(
                "<table><thead><tr><th>Month</th><th>Defects</th><th>MTBF (h)</th></tr></thead>"
                f"<tbody>{rows}</tbody></table>"
            )

    html = f"""<!doctype html>
<html>
<head>
  <meta charset=\"utf-8\" />
  <title>{spec.client} - {spec.job_type} - {spec.year}</title>
  <style>
    body {{ font-family: Georgia, serif; margin: 0; color: #1f2a44; }}
    .letterhead {{ background: #0b3251; color: #fff; padding: 18px 32px; font-weight: 700; letter-spacing: 0.08em; }}
    .title-page {{ padding: 42px 40px 28px; border-bottom: 3px solid #e5e7eb; }}
    .title-page h1 {{ margin: 0 0 12px; font-size: 34px; font-weight: 900; }}
    .meta strong {{ display: inline-block; min-width: 170px; }}
    .content {{ padding: 22px 34px 44px; }}
    h2 {{ font-size: 26px; border-left: 6px solid #c0392b; padding-left: 10px; margin-top: 30px; }}
    p {{ line-height: 1.6; }}
    .chart {{ width: 100%; max-width: 760px; border: 1px solid #ddd; }}
    table {{ border-collapse: collapse; width: 100%; margin-top: 12px; }}
    th, td {{ border: 1px solid #cbd5e1; padding: 8px; text-align: left; }}
    th {{ background: #eff6ff; }}
    footer {{ margin-top: 30px; font-size: 12px; color: #667085; border-top: 1px solid #ddd; padding-top: 8px; }}
  </style>
</head>
<body>
  <div class=\"letterhead\">ALINTA CONSULTING SERVICES</div>
  <section class=\"title-page\">
    <h1>TECHNICAL CONSULTING REPORT</h1>
    <div class=\"meta\">
      <p><strong>Client Name:</strong> {spec.client}</p>
      <p><strong>Report ID:</strong> {spec.report_id}</p>
      <p><strong>Report Date:</strong> {spec.year}-10-15</p>
      <p><strong>Client Contact:</strong> {spec.contact}</p>
      <p><strong>Author:</strong> {spec.author}</p>
      <p><strong>Report Type:</strong> {spec.job_type}</p>
      <p><strong>Site / Asset:</strong> {spec.site} / {spec.asset}</p>
    </div>
  </section>
  <main class=\"content\">
    {''.join(section_html)}
    <footer>Confidential - Consultancy Technical Report</footer>
  </main>
</body>
</html>
"""
    out.write_text(html, encoding="utf-8")


def header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(20 * mm, A4[1] - 14 * mm, "ALINTA CONSULTING SERVICES")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(A4[0] - 20 * mm, 10 * mm, f"Page {doc.page}")
    canvas.drawString(20 * mm, 10 * mm, "Confidential - Consultancy Technical Report")
    canvas.restoreState()


def write_pdf(out: Path, spec: ReportSpec, sections: dict[str, list[str]], chart_path: Path, months: list[str], defects: list[int], mtbf: list[float]) -> None:
    styles = getSampleStyleSheet()
    heading = ParagraphStyle("H", parent=styles["Heading2"], fontSize=16, spaceAfter=8, textColor=colors.HexColor("#0b3251"))
    normal = ParagraphStyle("N", parent=styles["BodyText"], fontSize=10, leading=14)

    story = []
    story.append(Paragraph("TECHNICAL CONSULTING REPORT", styles["Title"]))
    story.append(Spacer(1, 10))
    meta_lines = [
        f"<b>Client Name:</b> {spec.client}",
        f"<b>Report ID:</b> {spec.report_id}",
        f"<b>Report Date:</b> {spec.year}-10-15",
        f"<b>Client Contact:</b> {spec.contact}",
        f"<b>Author:</b> {spec.author}",
        f"<b>Report Type:</b> {spec.job_type}",
        f"<b>Site / Asset:</b> {spec.site} / {spec.asset}",
    ]
    for line in meta_lines:
        story.append(Paragraph(line, normal))
        story.append(Spacer(1, 3))

    story.append(PageBreak())

    for name, paras in sections.items():
        story.append(Paragraph(name, heading))
        for p in paras:
            story.append(Paragraph(p, normal))
            story.append(Spacer(1, 8))

        if name == "Results":
            story.append(Paragraph("Results Trend Chart", heading))
            story.append(RLImage(str(chart_path), width=170 * mm, height=80 * mm))
            story.append(Spacer(1, 8))
            data = [["Month", "Defects", "MTBF (h)"]] + [[m, str(d), str(t)] for m, d, t in zip(months, defects, mtbf)]
            table = Table(data, colWidths=[35 * mm, 35 * mm, 35 * mm])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6f0ff")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 10))

    doc = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=20 * mm, bottomMargin=16 * mm)
    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def main() -> None:
    out_dir = Path("sample_docs")
    out_dir.mkdir(parents=True, exist_ok=True)
    chart_dir = out_dir / "charts"
    chart_dir.mkdir(parents=True, exist_ok=True)

    specs = [
        ReportSpec(2016, "NorthRiver Energy", "Piping Integrity Assessment", "pdf", "Unit 1", "Line 3 / Weld Cluster A", "NRE-PIA-2016-011", "M. Sullivan", "D. Harding"),
        ReportSpec(2017, "BlueMesa Utilities", "Electrical Reliability Audit", "docx", "Substation East", "Switchboard SB-2", "BMU-ERA-2017-024", "C. Tran", "R. Bell"),
        ReportSpec(2018, "HarborStone Manufacturing", "Mechanical Overhaul Review", "html", "Process Hall B", "Compressor Train C", "HSM-MOR-2018-039", "A. Keane", "L. Porter"),
        ReportSpec(2019, "NorthRiver Energy", "Civil Structural Condition Survey", "pdf", "Tank Farm", "Pipe Rack PR-7", "NRE-CSS-2019-053", "M. Sullivan", "D. Harding"),
        ReportSpec(2020, "BlueMesa Utilities", "Piping Integrity Assessment", "docx", "Water Treatment Unit", "Slurry Line SL-9", "BMU-PIA-2020-067", "C. Tran", "R. Bell"),
        ReportSpec(2021, "HarborStone Manufacturing", "Electrical Reliability Audit", "html", "Utility Corridor", "MCC-4", "HSM-ERA-2021-081", "A. Keane", "L. Porter"),
        ReportSpec(2022, "NorthRiver Energy", "Mechanical Overhaul Review", "pdf", "Gas Compression Area", "Pump P-214", "NRE-MOR-2022-099", "M. Sullivan", "D. Harding"),
        ReportSpec(2023, "BlueMesa Utilities", "Civil Structural Condition Survey", "html", "Reservoir Access Bridge", "Span S2", "BMU-CSS-2023-113", "C. Tran", "R. Bell"),
        ReportSpec(2024, "HarborStone Manufacturing", "Piping Integrity Assessment", "docx", "Finishing Line", "Steam Header SH-1", "HSM-PIA-2024-129", "A. Keane", "L. Porter"),
        ReportSpec(2025, "NorthRiver Energy", "Electrical Reliability Audit", "pdf", "Main Distribution Room", "Relay Panel PRP-5", "NRE-ERA-2025-141", "M. Sullivan", "D. Harding"),
    ]

    created = []
    for i, spec in enumerate(specs, start=1):
        sections = section_paragraphs(spec)
        chart_name = f"report_v2_{i:02d}_{spec.year}_{spec.client.lower().replace(' ', '_')}_trend.png"
        chart_path = chart_dir / chart_name
        months, defects, mtbf = create_chart(
            chart_path,
            title=f"{spec.client} - {spec.job_type} ({spec.year})",
            seed=i + spec.year,
        )

        filename = f"report_v2_{i:02d}_{spec.year}_{spec.client.lower().replace(' ', '_')}_{spec.job_type.lower().replace(' ', '_')}.{spec.fmt}"
        out = out_dir / filename

        if spec.fmt == "docx":
            write_docx(out, spec, sections, chart_path, months, defects, mtbf)
        elif spec.fmt == "html":
            write_html(out, spec, sections, chart_path, months, defects, mtbf)
        else:
            write_pdf(out, spec, sections, chart_path, months, defects, mtbf)

        created.append(out.name)

    print("CREATED_RICH_REPORTS")
    for name in created:
        print(name)


if __name__ == "__main__":
    main()
