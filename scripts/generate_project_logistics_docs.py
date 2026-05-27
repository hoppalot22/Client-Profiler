from __future__ import annotations

import argparse
import csv
import json
import sys
from dataclasses import dataclass
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from openpyxl import Workbook

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from client_profiler import ClientProfiler, ProfilerConfig


SAMPLE_DIR = PROJECT_ROOT / "sample_docs"
LOGISTICS_DIR = SAMPLE_DIR / "project_logistics"

REPORT_TYPE_LABELS = {
    "piping_integrity_assessment": "Piping Integrity Assessment",
    "electrical_reliability_audit": "Electrical Reliability Audit",
    "mechanical_overhaul_review": "Mechanical Overhaul Review",
    "civil_structural_condition_survey": "Civil Structural Condition Survey",
}

CLIENT_SITES = {
    "Northriver Energy": "Kestrel Gas Plant",
    "Bluemesa Utilities": "Dry Creek Power Station",
    "Harborstone Manufacturing": "Port Melville Works",
}

CONTRACTORS = [
    "Vector Field Services",
    "Axis Shutdown Support",
    "Crestline Site Access",
    "Meridian Commissioning Partners",
]

WORKERS = [
    "Alyssa Tran",
    "Jordan Pike",
    "Marco Singh",
    "Rebecca Holt",
    "Nathan Wu",
    "Talia Mercer",
]


@dataclass
class ReportDescriptor:
    source_path: Path
    stem: str
    version_tag: str
    sequence: int
    year: int
    client_name: str
    report_type_slug: str
    report_type_label: str
    project_name: str
    project_code: str
    site_name: str


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate realistic logistics docs for sample reports.")
    parser.add_argument("--output-dir", default=str(LOGISTICS_DIR), help="Directory to write generated documents into.")
    parser.add_argument("--ingest", action="store_true", help="Ingest generated documents after creating them.")
    parser.add_argument("--force", action="store_true", help="Force re-ingestion when used with --ingest.")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    descriptors = _collect_reports(SAMPLE_DIR)
    manifest = _generate_documents(descriptors, output_dir)
    print(json.dumps({"generated_documents": len(manifest), "output_dir": str(output_dir)}, indent=2))

    if args.ingest:
        summary = _ingest_and_verify(output_dir, manifest, args.force)
        print(json.dumps(summary, indent=2))


def _collect_reports(sample_dir: Path) -> list[ReportDescriptor]:
    descriptors: list[ReportDescriptor] = []
    for path in sorted(sample_dir.iterdir()):
        if not path.is_file():
            continue
        if not path.stem.startswith("report"):
            continue
        descriptor = _parse_report_descriptor(path)
        if descriptor:
            descriptors.append(descriptor)
    return descriptors


def _parse_report_descriptor(path: Path) -> ReportDescriptor | None:
    parts = path.stem.split("_")
    if len(parts) < 6 or parts[0] != "report":
        return None

    idx = 1
    version_tag = "base"
    if parts[idx].startswith("v"):
        version_tag = parts[idx]
        idx += 1

    try:
        sequence = int(parts[idx])
        year = int(parts[idx + 1])
    except (ValueError, IndexError):
        return None

    remainder = parts[idx + 2 :]
    report_type_slug = _find_report_type_slug(remainder)
    if not report_type_slug:
        return None
    report_type_parts = report_type_slug.split("_")
    client_parts = remainder[: len(remainder) - len(report_type_parts)]
    if not client_parts:
        return None

    client_name = " ".join(client_parts).title()
    report_type_label = REPORT_TYPE_LABELS[report_type_slug]
    project_name = report_type_label
    project_code = _derive_project_code(client_name, project_name, year)

    return ReportDescriptor(
        source_path=path,
        stem=path.stem,
        version_tag=version_tag,
        sequence=sequence,
        year=year,
        client_name=client_name,
        report_type_slug=report_type_slug,
        report_type_label=report_type_label,
        project_name=project_name,
        project_code=project_code,
        site_name=CLIENT_SITES.get(client_name, "Main Process Facility"),
    )


def _find_report_type_slug(parts: list[str]) -> str | None:
    joined = "_".join(parts)
    for slug in REPORT_TYPE_LABELS:
        if joined.endswith(slug):
            return slug
    return None


def _derive_project_code(client_name: str, project_name: str, year: int) -> str:
    client = "".join(word[0] for word in client_name.split()[:3]).upper()
    words = [word[:3].upper() for word in project_name.split()[:3]]
    return "-".join([client, *words, str(year)])


def _generate_documents(descriptors: list[ReportDescriptor], output_dir: Path) -> list[dict[str, Any]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: list[dict[str, Any]] = []

    quote_formats = [".txt", ".html", ".docx"]
    variation_formats = [".md", ".html", ".docx"]
    po_formats = [".xlsx", ".csv", ".txt"]
    email_formats = [".md", ".html", ".txt"]
    access_formats = [".docx", ".html", ".csv"]

    for descriptor in descriptors:
        folder = output_dir / descriptor.stem
        folder.mkdir(parents=True, exist_ok=True)
        doc_date = date(descriptor.year, ((descriptor.sequence - 1) % 12) + 1, min(24, 6 + descriptor.sequence))
        contractor = CONTRACTORS[(descriptor.sequence - 1) % len(CONTRACTORS)]
        crew = [WORKERS[(descriptor.sequence + offset) % len(WORKERS)] for offset in range(3)]

        quote_number = f"Q-{descriptor.year}-{descriptor.sequence:02d}{'B' if descriptor.version_tag != 'base' else 'A'}"
        variation_number = f"{quote_number}-REV1"
        po_number = f"PO-{descriptor.year}-{descriptor.sequence:02d}{'2' if descriptor.version_tag != 'base' else '1'}"
        access_ref = f"AR-{descriptor.year}-{descriptor.sequence:02d}{'R' if descriptor.version_tag != 'base' else 'P'}"

        manifest.extend(
            [
                _write_quote(folder, descriptor, doc_date, contractor, quote_number, quote_formats[(descriptor.sequence - 1) % len(quote_formats)]),
                _write_quote_variation(folder, descriptor, doc_date + timedelta(days=2), contractor, variation_number, quote_number, variation_formats[(descriptor.sequence - 1) % len(variation_formats)]),
                _write_purchase_order(folder, descriptor, doc_date + timedelta(days=5), po_number, quote_number, po_formats[(descriptor.sequence - 1) % len(po_formats)]),
                _write_email_chain(folder, descriptor, doc_date + timedelta(days=7), contractor, quote_number, po_number, crew, email_formats[(descriptor.sequence - 1) % len(email_formats)]),
                _write_access_pack(folder, descriptor, doc_date + timedelta(days=9), contractor, access_ref, po_number, crew, access_formats[(descriptor.sequence - 1) % len(access_formats)]),
            ]
        )

    return manifest


def _write_quote(folder: Path, descriptor: ReportDescriptor, doc_date: date, contractor: str, quote_number: str, suffix: str) -> dict[str, Any]:
    filename = folder / f"quote_{quote_number.lower().replace('-', '_')}{suffix}"
    title = f"Budget Estimate for {descriptor.project_name}"
    body = {
        "Client": descriptor.client_name,
        "Project": descriptor.project_name,
        "Project Code": descriptor.project_code,
        "Quote Number": quote_number,
        "Date": doc_date.isoformat(),
        "Site": descriptor.site_name,
        "Supplier": contractor,
        "Linked Report": descriptor.source_path.name,
        "Scope": f"Mobilise specialist crew for {descriptor.project_name.lower()} including shutdown planning, testing attendance, and reporting close-out.",
        "Commercial Notes": "Pricing includes site induction, temporary access equipment, and after-hours standby coverage.",
    }
    _write_by_suffix(filename, title, body)
    return {"path": str(filename), "expected_kind": "quote", "project_name": descriptor.project_name}


def _write_quote_variation(
    folder: Path,
    descriptor: ReportDescriptor,
    doc_date: date,
    contractor: str,
    variation_number: str,
    base_quote_number: str,
    suffix: str,
) -> dict[str, Any]:
    filename = folder / f"quote_variation_{variation_number.lower().replace('-', '_')}{suffix}"
    title = f"Variation Request {variation_number}"
    body = {
        "Client": descriptor.client_name,
        "Project": descriptor.project_name,
        "Project Code": descriptor.project_code,
        "Quote Number": variation_number,
        "Reference": base_quote_number,
        "Date": doc_date.isoformat(),
        "Reason": "Additional cable isolation point verification and standby labour requested after shutdown sequence review.",
        "Change": "Add one electrical test technician for the night shift and extend elevated work platform hire by one day.",
        "Approval Path": "Submit revised quotation to site maintenance planner for PO uplift.",
    }
    _write_by_suffix(filename, title, body)
    return {"path": str(filename), "expected_kind": "quote", "project_name": descriptor.project_name}


def _write_purchase_order(
    folder: Path,
    descriptor: ReportDescriptor,
    doc_date: date,
    po_number: str,
    quote_number: str,
    suffix: str,
) -> dict[str, Any]:
    filename = folder / f"purchase_order_{po_number.lower().replace('-', '_')}{suffix}"
    body = {
        "Client": descriptor.client_name,
        "Project": descriptor.project_name,
        "Project Code": descriptor.project_code,
        "Purchase Order Number": po_number,
        "Date": doc_date.isoformat(),
        "Quote Reference": quote_number,
        "Service Window": f"{doc_date.isoformat()} to {(doc_date + timedelta(days=3)).isoformat()}",
        "Cost Centre": f"{descriptor.project_code}-OPS",
        "Description": f"Approved field services and logistics support for {descriptor.project_name.lower()} at {descriptor.site_name}.",
    }
    _write_by_suffix(filename, "Purchase Order", body)
    return {"path": str(filename), "expected_kind": "purchase_order", "project_name": descriptor.project_name}


def _write_email_chain(
    folder: Path,
    descriptor: ReportDescriptor,
    doc_date: date,
    contractor: str,
    quote_number: str,
    po_number: str,
    crew: list[str],
    suffix: str,
) -> dict[str, Any]:
    filename = folder / f"email_chain_{descriptor.project_code.lower().replace('-', '_')}{suffix}"
    subject = f"RE: {descriptor.project_name} mobilisation and access hold points"
    chain = [
        {
            "From": f"maintenance.planner@{descriptor.client_name.split()[0].lower()}.example.com",
            "To": f"coordinator@{contractor.split()[0].lower()}.example.com",
            "Sent": f"{doc_date.isoformat()} 07:42",
            "Subject": subject,
            "Body": f"Please confirm the crew list for {descriptor.project_name}, linked to quote {quote_number} and PO {po_number}. Shutdown gate opens at 05:30 on {(doc_date + timedelta(days=4)).isoformat()}.",
        },
        {
            "From": f"coordinator@{contractor.split()[0].lower()}.example.com",
            "To": f"maintenance.planner@{descriptor.client_name.split()[0].lower()}.example.com",
            "Sent": f"{doc_date.isoformat()} 09:18",
            "Subject": subject,
            "Body": f"Crew nominated: {', '.join(crew)}. All hold current inductions; we still require confined-space approval for the LV room entry and confirmation of escort arrangements.",
        },
        {
            "From": f"shutdown.supervisor@{descriptor.client_name.split()[0].lower()}.example.com",
            "To": f"coordinator@{contractor.split()[0].lower()}.example.com",
            "Sent": f"{(doc_date + timedelta(days=1)).isoformat()} 14:05",
            "Subject": subject,
            "Body": f"Escort approved for day shift only. Night shift access will be released once the permit office has the final SWMS and the elevated work platform registration.",
        },
    ]
    payload = {
        "Client": descriptor.client_name,
        "Project": descriptor.project_name,
        "Project Code": descriptor.project_code,
        "Date": doc_date.isoformat(),
        "Quote Number": quote_number,
        "Purchase Order Number": po_number,
        "Email Chain": chain,
    }
    _write_by_suffix(filename, subject, payload, rich=True)
    return {"path": str(filename), "expected_kind": "email_chain", "project_name": descriptor.project_name}


def _write_access_pack(
    folder: Path,
    descriptor: ReportDescriptor,
    doc_date: date,
    contractor: str,
    access_ref: str,
    po_number: str,
    crew: list[str],
    suffix: str,
) -> dict[str, Any]:
    filename = folder / f"access_request_{access_ref.lower().replace('-', '_')}{suffix}"
    payload = {
        "Client": descriptor.client_name,
        "Project": descriptor.project_name,
        "Project Code": descriptor.project_code,
        "Access Request Number": access_ref,
        "Purchase Order Number": po_number,
        "Date": doc_date.isoformat(),
        "Requested By": contractor,
        "Site": descriptor.site_name,
        "Workers": ", ".join(crew),
        "Request": "Issue temporary access, plant induction confirmation, and permit office approval for shutdown execution window.",
        "Approval": f"Approved by {descriptor.client_name} permit office subject to daily pre-start and escorted entry for LV room isolation checks.",
    }
    _write_by_suffix(filename, "Access Request and Approval", payload)
    return {"path": str(filename), "expected_kind": "access_request", "project_name": descriptor.project_name}


def _write_by_suffix(path: Path, title: str, body: dict[str, Any], rich: bool = False) -> None:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        lines = [title, "=" * len(title), ""]
        for key, value in body.items():
            lines.append(f"{key}: {_render_value(value)}")
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    if suffix in {".md", ".markdown"}:
        lines = [f"# {title}", ""]
        for key, value in body.items():
            if isinstance(value, list):
                lines.append(f"## {key}")
                lines.append("")
                for item in value:
                    lines.append(f"- {_render_value(item)}")
            else:
                lines.append(f"- **{key}:** {_render_value(value)}")
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    if suffix == ".html":
        rows = []
        for key, value in body.items():
            rows.append(f"<tr><th>{key}</th><td>{_render_html_value(value)}</td></tr>")
        html = f"""<!doctype html>
<html lang=\"en\">
<head><meta charset=\"utf-8\" /><title>{title}</title></head>
<body>
  <main>
    <h1>{title}</h1>
    <table>{''.join(rows)}</table>
  </main>
</body>
</html>
"""
        path.write_text(html, encoding="utf-8")
        return

    if suffix == ".docx":
        doc = DocxDocument()
        doc.add_heading(title, level=1)
        for key, value in body.items():
            doc.add_paragraph(f"{key}: {_render_value(value)}")
        doc.save(path)
        return

    if suffix == ".csv":
        with path.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["Field", "Value"])
            for key, value in body.items():
                writer.writerow([key, _render_value(value)])
        return

    if suffix == ".xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = "Details"
        ws.append([title, ""])
        ws.append(["Field", "Value"])
        for key, value in body.items():
            ws.append([key, _render_value(value)])
        wb.save(path)
        return

    raise ValueError(f"Unsupported suffix for generation: {suffix}")


def _render_value(value: Any) -> str:
    if isinstance(value, list):
        return "; ".join(_render_value(item) for item in value)
    if isinstance(value, dict):
        return " | ".join(f"{key}={_render_value(item)}" for key, item in value.items())
    return str(value)


def _render_html_value(value: Any) -> str:
    if isinstance(value, list):
        items = "".join(f"<li>{_render_html_value(item)}</li>" for item in value)
        return f"<ul>{items}</ul>"
    if isinstance(value, dict):
        items = "".join(f"<li><strong>{key}</strong>: {_render_html_value(item)}</li>" for key, item in value.items())
        return f"<ul>{items}</ul>"
    return str(value)


def _ingest_and_verify(output_dir: Path, manifest: list[dict[str, Any]], force: bool) -> dict[str, Any]:
    profiler = ClientProfiler(ProfilerConfig())
    results = profiler.ingest_directory(output_dir, force_reingest=force)

    expected_by_path = {str(Path(item["path"]).resolve()): item for item in manifest}
    checks = []
    mismatches = []
    for result in results:
        path = str(Path(result.get("path", "")).resolve())
        expected = expected_by_path.get(path)
        if not expected:
            continue
        actual_kind = result.get("document_kind")
        if actual_kind is None:
            record = profiler.storage.get_latest_document_record(result.get("path", ""))
            metadata = record.get("metadata", {}) if isinstance(record, dict) else {}
            if isinstance(metadata, dict):
                actual_kind = metadata.get("document_kind")
        ok = actual_kind == expected["expected_kind"]
        checks.append({
            "path": result.get("path"),
            "expected_kind": expected["expected_kind"],
            "actual_kind": actual_kind,
            "status": result.get("status"),
            "ok": ok,
        })
        if not ok:
            mismatches.append(checks[-1])

    return {
        "ingested": len(results),
        "verified": len(checks),
        "mismatches": mismatches,
    }


if __name__ == "__main__":
    main()